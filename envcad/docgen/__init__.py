# -*- coding: utf-8 -*-
"""文档自动化层：从知识层/设计层一键生成工程文档。

生成器：
  spec_doc    《结构设计总说明》DOCX
  calc_book   《结构计算书》DOCX
  bom_xlsx    《材料表》XLSX

统一的汉字排版 helper：把中英文都设成指定字体（含 w:eastAsia），
避免 Word 里中文掉成默认字体。
"""
from __future__ import annotations

# python-docx 改为函数内惰性导入：XLSX 类命令（openpyxl）不再连带加载 docx。
from typing import TYPE_CHECKING

if TYPE_CHECKING:  # 仅供类型标注，运行时不导入
    from docx import Document


def _set_run_font(run, name: str = "宋体", size: int = None,
                  bold: bool = False, color=None):
    """设置 run 字体（中英文字体一致，含 eastAsia）。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), name)


def new_cn_doc(title: str = None, title_size: int = 18) -> "Document":
    """新建中文文档，Normal 样式锁定宋体，标题用黑体。"""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), "宋体")
    if title:
        h = doc.add_heading(level=0)
        run = h.add_run(title)
        _set_run_font(run, "黑体", title_size, bold=True)
    return doc


def add_heading_cn(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    _set_run_font(run, "黑体", 14 - level, bold=True)
    return h


def add_para_cn(doc: Document, text: str, size: int = 10.5, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, "宋体", size, bold)
    return p


def add_table_cn(doc: Document, header, rows, col_widths=None):
    """中文表格，表头加粗，单元格锁定宋体。"""
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        run = hdr[i].paragraphs[0].add_run(str(htext))
        _set_run_font(run, "黑体", 10.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            _set_run_font(run, "宋体", 10.5)
    return t


__all__ = [
    "new_cn_doc", "add_heading_cn", "add_para_cn", "add_table_cn", "_set_run_font",
]
